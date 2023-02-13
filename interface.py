# for infopage
from main import inPathSendableObj, outPathSendable, inPathSendableSub, typeQ
from docx import Document
from pathlib import Path
from os import path
from win32api import GetSystemMetrics
# from tkinter import *
# Explicit imports to satisfy Flake8
from tkinter import Tk, Canvas, Entry, IntVar, Button, PhotoImage, END, messagebox as mg, StringVar
from tkinter.ttk import Combobox
from tkinter.filedialog import askdirectory, askopenfile
from random import sample
# completed importing stuffs.
OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(rf"{OUTPUT_PATH}\assets\frame0")


# temp code
# outPathSendable = "OutDoc.doc"
# inPathSendableObj = "C:/Users/LENOVO/OneDrive/Desktop/QuesLong..docx"
# inPathSendableSub = "C:/Users/LENOVO/OneDrive/Desktop/QuesLong..docx"
# typeQ ="Subjective"
# Universally used 

class Dummy():
    
    paragraphs = "N" 
    def __init__(self) -> None:
        pass
    def add_paragraph(self,text):
        pass
    def add_heading(self,text):
        pass

letters = ["A", "B", "C", "D", "E"]
try:
    if typeQ=="Objective":
        InDocObj = Document(inPathSendableObj)
        InDocSub = Dummy()
    elif typeQ=="Subjective":
        InDocObj = Dummy()
        InDocSub = Document(inPathSendableSub)
    else:
        InDocObj = Document(inPathSendableObj)
        InDocSub = Document(inPathSendableSub)

    noPara_ObjFile = len(InDocObj.paragraphs)
    noPara_SubFile = len(InDocSub.paragraphs)
except Exception as e:
    mg.showerror("Path Error Ocurred", "Please Restart the Application.")
    quit()
noOfQuestFound_ObjFile = noPara_ObjFile//5
noOfQuestFound_SubFile = noPara_SubFile
OutDoc = None
# getting system Metrices
width = GetSystemMetrics(0)
height = GetSystemMetrics(1)

# for making graph more imaginable
Y_CONST = 200
# returns a relative path of files


def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

# function which runs the query of selecting path


def pathSelect(no):
    # 1 equivalents to searching for a file
    # else searching for a folder
    if no == 1 or no == 2:
        global inPathSendableObj, inPathSendableSub, outPathSendable
        try:
            # getting the file path as entered picked by user
            # for file selection
            path = askopenfile().name
            # lenght of the path string
            LenPath = len(path)-1
            # converting text to lowercase for validation and checking wheather the input file is input is or not
            if (path[LenPath:LenPath-4:-1].lower() != "xcod" and path[LenPath:LenPath-3:-1].lower() != "cod"):
                mg.showerror("File Selection Error",
                             "The file you selected is not a word file.")
                return
            # this is done to change the value of Input folder path when user decides to change the previous file
            global InDocObj, InDocSub, noPara_ObjFile, noPara_SubFile, noOfQuestFound_SubFile, noOfQuestFound_ObjFile
            try:
                if no == 1:
                    inPathSendableObj = path
                    # redefining instance of Document class
                    InDocObj = Document(inPathSendableObj)
                    ObjFileSec.delete(0, END)
                    ObjFileSec.insert(0, path)
                    ObjFileSec.xview_moveto(1)
                    canvas.itemconfig(
                        changable1, text=f"Question Found: {len(InDocObj.paragraphs)//5}")
                    # Redefining the varaible that depends on the InDocObj
                    noPara_ObjFile = len(InDocObj.paragraphs)
                    noOfQuestFound_ObjFile = noPara_ObjFile//5

                else:
                    inPathSendableSub = path
                    # redefining instance of Document class
                    global InDocSub
                    InDocSub = Document(inPathSendableSub)
                    SubFileSec.delete(0, END)
                    SubFileSec.insert(0, path)
                    SubFileSec.xview_moveto(1)
                    canvas.itemconfig(
                        changable2, text=f"Question Found: {len(InDocSub.paragraphs)}")
                    # Redefining the variable that depends on InDocSub
                    noPara_SubFile = len(InDocSub.paragraphs)
                    noOfQuestFound_SubFile = noPara_SubFile
            except:
                mg.showerror(
                    "OCCURRENCE OF ERROR", "Please check the Path provided or Re-Start the Program.")
            # moves the cursor to right most part
        # Global exception handler
        except Exception as e:
            mg.showerror("Error", "Input File is not selected.")
            ObjFileSec.insert(0, "")
    else:
        # For output path selection
        path = askdirectory()
        outPathSendable = path
        # deletes what's already their
        OutEntry.delete(0, END)
        OutEntry.insert(0, path)
        OutEntry.xview_moveto(1)

# this function randomizes the option in question set


def optionRandomizer(OutDoc, paraNum):
    '''This function randomizes the option ie changes the order or option so that they won't be same'''
    # creating a random list in range 1-4
    default_list = sample(range(1, 5), 4)
    # since no of option is always 4
    for index, val in enumerate(default_list):
        # text to be appended in OutDoc
        reqText = f"{letters[index]}) {InDocObj.paragraphs[paraNum+val if (paraNum+val)<noOfQuestFound_ObjFile*5 else paraNum].text}"
        OutDoc.add_paragraph(reqText)

# this code creates random question for subjective question


def SUB_QUES(NoOfQues):
    ques_pos_list = None
    """
    This function generates random question for subjective question set
    """
    ques_pos_list = sample(range(0, noPara_SubFile),
                           NoOfQues)  # NO_QUESTION_SUB
    for counter, ranNo in enumerate(ques_pos_list):
        # ranNo is the random number
        # Getting text from input file
        paraCode = InDocSub.paragraphs[ranNo]
        # getting the actual code
        InpQues = paraCode.text
        # copying that text into new file
        OutDoc.add_paragraph(f"{counter+1}) {InpQues}")


def OBJ_QUES(NoOfQues):
    ques_pos_list = None
    # list which acts as a template for question fetching
    ques_pos_list = sample(range(0, noPara_ObjFile, 5), NoOfQues)
    for counter, ranNo in enumerate(ques_pos_list):
        # ranNo is the random number
        # Getting text from input file
        paraCode = InDocObj.paragraphs[ranNo]
        InpQues = paraCode.text
        # copying that text into new file
        OutDoc.add_paragraph(f"{counter+1}) {InpQues}")
        # function to randomize options
        optionRandomizer(OutDoc, ranNo)

# -----------------------------------------------------------------
# function which does the main processing


def PROCESSOR(*args):
    """
        Basic Parameters Structure:
        0 - input File Objective
        1-output File
        2-no of required question Objective
        3-no of required question Subjective
        4-no of Papers required
        5-input File Subjective
    """
    global typeQ, OutDoc  # type of  question user required
    typeQ = OPTION_VAR.get()

    # Getting no of paragraphs
    global noPara_ObjFile, noPara_SubFile
    if args[4]==0:
        mg.showerror("NO CREATION ERROR",
                            "Number of Paper is 0.Must be Greater than 0.")
        return
    # papNo -> No of paper required
    for papNo in range(args[4]):
        # creates new instances which creates new file
        OutDoc = Document()
        # checking the type of type user wants
        if typeQ == "Subjective":
            # doesn't check when the index in 3 ie doesn't check the objective wheather it is zero or not 
            if any(arg == "" or arg == 0 for ind,arg in enumerate(args) if ind!=2 and ind!=0):
                mg.showerror("Text Field Blank Error",
                            "Please Enter no of Subjective Question required.")
                return
            # code below adds subjective question to file
            if args[3] > (noOfQuestFound_SubFile):  
                mg.showerror("Question Deficit Error",
                            "Number of Question required exceeded Number of Question provided.")
                return
            OutDoc.add_heading(f"Answer the following Question", 2)
            SUB_QUES(NO_QUESTION_SUB.get())
        # When Subjective only
        elif typeQ == "Objective":
            if any(arg == "" or arg == 0 for ind,arg in enumerate(args) if ind!=3 and ind!=5):
                mg.showerror("Text Field Blank Error",
                            "Please enter number of Objective Question required.")
                return
            if args[2] > (noOfQuestFound_ObjFile):  
                mg.showerror("Question Deficit Error",
                            "Number of Question required exceeded Number of Question provided.")
                return
            OutDoc.add_heading(f"Group A :Multilpe Choice Question {NO_QUESTION_OBJ.get()}x1={NO_QUESTION_OBJ.get()}",2)
            OBJ_QUES(NO_QUESTION_OBJ.get())
        #When both subjective and objective 
        elif typeQ == "Subjective and Objective":
            if any(arg == "" or arg == 0 for arg in args):
                mg.showerror("Text Field Blank Error",
                            "Required Information must be provided.")
                return
            if args[2] > (noOfQuestFound_ObjFile) or args[3] > (noOfQuestFound_SubFile):  
                mg.showerror("Question Deficit Error",
                            "Number of Question required exceeded Number of Question provided.")
                return
            OutDoc.add_heading(f"Group A :Multilpe Choice Question {NO_QUESTION_OBJ.get()}x1={NO_QUESTION_OBJ.get()}",2)
            OBJ_QUES(NO_QUESTION_OBJ.get())
            OutDoc.add_heading(f"Group B",2)
            SUB_QUES(NO_QUESTION_SUB.get())
        else:
            mg.showerror("TYPE SELECTION ERROR",
                        "TYPE SELECTED IS NOT RECOGNIZED.")
            # saves every new file as in instances
        OutDoc.save(path.join(outPathSendable, f"Doc{papNo}.doc"))

    mg.showinfo("Completion of Task",
                "Successfully Created File in"+outPathSendable)

# CREATES A NEW WINDOW
window = Tk()
# setting window name
window.title("DeQuestify")
# no_of question required
NO_QUESTION_OBJ = IntVar()
NO_QUESTION_SUB = IntVar()
# no of paper required
NO_PAPER_REQUIRED = IntVar()
OPTION_VAR = StringVar()
# sets geometry of window
window.geometry(f"1000x{height-90}+{(width//2)-500}+5")
window.configure(bg="#363740")
# creating a blackish color canvas
canvas = Canvas(
    window,
    bg="#363740",
    height=620,
    width=1000,
    bd=0,
    highlightthickness=0,
    relief="ridge"
)
canvas.place(x=0, y=0)
# creating main frame
canvas.create_rectangle(
    195.0,
    15.0,
    794.0,
    658.0,
    fill="#FFFFFF",
    outline="")

# CODE BLOCK STARTS -> SUBMISSION BUTTON PATH
GenerateButton = PhotoImage(
    file=relative_to_assets("button_1.png"))
BtnGene = Button(
    image=GenerateButton,
    borderwidth=0,
    highlightthickness=0,
    command=lambda: PROCESSOR(
        inPathSendableObj, outPathSendable,NO_QUESTION_OBJ.get(), NO_QUESTION_SUB.get(), NO_PAPER_REQUIRED.get(),inPathSendableSub),
    relief="flat"
)
BtnGene.place(
    x=437.0,
    y=555.0,
    width=126.0,
    height=48.0
)
# CODE BLOCK ENDS -> SUBMISSION BUTTON
# CODE BLOCK STARTS -> INPUT DIRECTROY OBJECTIIVE QUESTION PATH
ImgObjFile = PhotoImage(
    file=relative_to_assets("entry_1.png"))
entry_bg_1 = canvas.create_image(
    328.0,
    Y_CONST+51.0,
    image=ImgObjFile
)
ObjFileSec = Entry(
    bd=0,
    bg="#FCFDFE",
    fg="#000716",
    highlightthickness=0,

)
ObjFileSec.place(
    x=227.0,
    y=Y_CONST+33.0,
    width=150.0,
    height=35.0
)
ObjFileSec.insert(0, inPathSendableObj)
ObjFileSec.xview_moveto(1)

PathPickerObj = PhotoImage(file=ASSETS_PATH / "path_picker.png")
path_picker_button = Button(
    image=PathPickerObj,
    text='',
    compound='center',
    fg='white',
    borderwidth=0,
    highlightthickness=0,
    command=lambda: pathSelect(1),
    relief='flat')
path_picker_button.place(
    x=400, y=Y_CONST+39,
    width=24,
    height=22)

canvas.create_text(
    209.0,
    Y_CONST+5.0,
    anchor="nw",
    text="INPUT DIRECTORY (OBJECTIVE)",
    fill="#565863",
    font=("Mulish Bold", 14 * -1)
)

# CODE BLOCK ENDS -> INPUT DIRECTORY PATH FOR SUBJECTIVE QUESTION
PathPickerSub = PhotoImage(file=ASSETS_PATH / "path_picker.png")
path_picker_button = Button(
    image=PathPickerSub,
    text='',
    compound='center',
    fg='white',
    borderwidth=0,
    highlightthickness=0,
    command=lambda: pathSelect(2),
    relief='flat')
path_picker_button.place(
    x=400, y=Y_CONST+124,
    width=24,
    height=22)
ImgEntrySub = PhotoImage(
    file=relative_to_assets("entry_1.png"))
entry_bg_3 = canvas.create_image(
    328.0,
    Y_CONST+136.0,
    image=ImgEntrySub
)
SubFileSec = Entry(
    bd=0,
    bg="#FCFDFE",
    fg="#000716",
    highlightthickness=0,

)
SubFileSec.place(
    x=227.0,
    y=Y_CONST+118.0,
    width=150.0,
    height=35.0
)
SubFileSec.insert(0, inPathSendableSub)
SubFileSec.xview_moveto(1)

canvas.create_text(
    209.0,
    Y_CONST+90.0,
    anchor="nw",
    text="INPUT DIRECTORY (Subjective)",
    fill="#565863",
    font=("Mulish Bold", 14 * -1)
)
# CODE BLOCK ENDS -> INPUT DIRECTORY PATH
# CODE BLOCK STARTS -> OUTPUT DIRECTORY PATH
OutEntryImg = PhotoImage(
    file=relative_to_assets("entry_1.png"))
entry_bg_2 = canvas.create_image(
    328.0,
    Y_CONST+223.0,
    image=OutEntryImg
)
OutEntry = Entry(
    bd=0,
    bg="#FCFDFE",
    fg="#000716",
    highlightthickness=0
)
OutEntry.place(
    x=227.0,
    y=Y_CONST+205.0,
    width=150.0,
    height=35.0
)
OutEntry.insert(0, outPathSendable)
OutEntry.xview_moveto(1)

OutPath_Picker = PhotoImage(file=ASSETS_PATH / "path_picker.png")
path_picker_button2 = Button(
    image=OutPath_Picker,
    text='',
    compound='center',
    fg='white',
    borderwidth=0,
    highlightthickness=0,
    # executes when button is clicked
    command=lambda: pathSelect(3),
    relief='flat')
path_picker_button2.place(
    x=400, y=Y_CONST+210,
    width=24,
    height=22)
canvas.create_text(
    209.0,
    Y_CONST+175.0,
    anchor="nw",
    text="OUTPUT DIRECTORY",
    fill="#44454D",
    font=("Mulish Bold", 14 * -1)
)
# CODE BLOCK ENDS -> OUTPUT DIRECTORY PATH

# CODE BLOCK START -> NO OF REQUIRED QUESTION OBJECTIVE
ReqQuesObj_Img = PhotoImage(
    file=relative_to_assets("entry_2.png"))
entry_bg_3 = canvas.create_image(
    628.0,
    Y_CONST+51.0,
    image=ReqQuesObj_Img
)
ReqQuesEntryObj = Entry(
    bd=0,
    bg="#FCFDFE",
    fg="#000716",
    highlightthickness=0,
    textvariable=NO_QUESTION_OBJ
)
ReqQuesEntryObj.place(
    x=527.0,
    y=Y_CONST+33.0,
    width=202.0,
    height=35.0
)

canvas.create_text(
    519.0,
    Y_CONST+5.0,
    anchor="nw",
    text="REQUIRED QUESTIONS -> OBJECTIVE",
    fill="#44454C",
    font=("Mulish Bold", 14 * -1)
)
# CODE BLOCK END -> NO OF QUESTION REQUIRED OBJECTIVE
# CODE BLOCK START -> NO OF REQUIRED QUESTION SUBJECTIVE
ReqQuesSub_Img = PhotoImage(
    file=relative_to_assets("entry_1.png"))
entry_bg_3 = canvas.create_image(
    628.0,
    Y_CONST+136,
    image=ReqQuesSub_Img
)
ReqQuesEntrySub = Entry(
    bd=0,
    bg="#FCFDFE",
    fg="#000716",
    highlightthickness=0,
    textvariable=NO_QUESTION_SUB
)
ReqQuesEntrySub.place(
    x=527.0,
    y=Y_CONST+118,
    width=202.0,
    height=35.0
)

canvas.create_text(
    519.0,
    Y_CONST+90,
    anchor="nw",
    text="REQUIRED QUESTIONS -> SUBJECTIVE",
    fill="#44454C",
    font=("Mulish Bold", 14 * -1)
)
# CODE BLOCK END -> NO OF QUESTION REQUIRED SUBJECTIVE
# CODE BLOCK START -> NO OF FILE OR NO OF PAPER
NoPapReq = PhotoImage(
    file=relative_to_assets("entry_2.png"))
entry_bg_4 = canvas.create_image(
    628.0,
    Y_CONST+223,
    image=NoPapReq
)
NoReqPapEntry = Entry(
    bd=0,
    bg="#FCFDFE",
    fg="#000716",
    highlightthickness=0,
    textvariable=NO_PAPER_REQUIRED
)
NoReqPapEntry.place(
    x=527.0,
    y=Y_CONST+205.0,
    width=202.0,
    height=35.0
)

canvas.create_text(
    519.0,
    Y_CONST+175,
    anchor="nw",
    text="No of papers",
    fill="#44454D",
    font=("Mulish Bold", 14 * -1)
)


# Logo
canvas.create_text(
    288.0,
    40.0,
    anchor="nw",
    text="DeQuestify",
    fill="#A4A6B3",
    font=("Mulish Bold", 19 * -1)
)
# Image for logo of dequestify
image_image_1 = PhotoImage(
    file=relative_to_assets("image_1.png"))
image_1 = canvas.create_image(
    258.0,
    50.0,
    image=image_image_1
)

# Admin Access Text--
canvas.create_text(
    410.0,
    84.0,
    anchor="nw",
    text="Admin Access",
    fill="#252733",
    font=("Mulish Bold", 24 * -1)
)
# Question Found Box-------
canvas.create_rectangle(
    303.5,
    124.0,
    472.0,
    164.0,
    fill="#00B2FF",
    outline="")
canvas.create_rectangle(
    510.0,
    124.0,
    678.0,
    164.0,
    fill="#0258f7",
    outline="")   
# Question found Text
changable1 = canvas.create_text(
    327.0,
    133.0,
    tags="notext1",
    anchor="nw",
    text=f"Question Found :{len(InDocObj.paragraphs)//5}",
    fill="#FFFFFF",
    font=("Muli SemiBold", 16 * -1)
)
changable2 = canvas.create_text(
    527.0,
    133.0,
    tags="notext2",
    anchor="nw",
    text=f"Question Found :{len(InDocSub.paragraphs)}",
    fill="#FFFFFF",
    font=("Muli SemiBold", 16 * -1)
)
choices = ('Subjective', 'Objective', 'Subjective and Objective')
OPTION_VAR.set('Type of Question to be Generated - Any One')
# for selecting option
optionSelect = Combobox(window, values=choices, textvariable=OPTION_VAR)
optionSelect.config(width=40, font=("Arial", 14 * -1))
optionSelect.place(x=333, y=482.5)
if (typeQ=='Subjective'):optionSelect.current(0)
elif (typeQ=="Objective"):optionSelect.current(1)
else:optionSelect.current(2)
    
window.resizable(False, False)
window.mainloop()
