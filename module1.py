# import docx NOT python-docx
from docx import *
import random
# create an instance of a word document
InDoc = Document("C:\Users\LENOVO\OneDrive\Desktop\QuesLong..docx")
OutDoc = Document("OutDoc.doc")

# let n be the no of question user wants
n = 30
font = "Arial"
size=30 
bold = True
italic  = True

# Getting no of paragraphs
noPara = len(InDoc.paragraphs)
# actual len of Paragraphs
actNoPara = (noPara//5)*5
print(noPara//5,actNoPara)
letters = ["A","B","C","D","E"]


# function to calculate repeatition rate
def RepeatCalc(quesProd):
    '''
    This function takes no of question to be produced as arguement
    and check the repeatition rate on total no of question
    '''
    noOfQues = noPara//5
    noforQues= random.randrange(1,noOfQues)
    arrOfnumbers = []
    repeated=0
    for coun in range(quesProd):
        if(any(ints== noforQues for ints in arrOfnumbers)):
            repeated=repeated+1
        else:
            arrOfnumbers.append(noforQues)
    
    perRepeated = (repeated/quesProd)*100
    return perRepeated # rate of repeatition percentage

# randomizes the option 
def optionRandomizer(paraNum):
    '''This function randomizes the option ie changes the order or option so that they won't be same'''
    # creating a random list in range 1-4
    default_list = random.sample(range(1,5),4)
    # since no of option is always 4
    for index,val in enumerate(default_list):
        #text to be appended in OutDoc
        reqText = f"{letters[index]}) {InDoc.paragraphs[paraNum+val].text}"
        OutDoc.add_paragraph(reqText)

#random no checker

#list which acts as a template for question fetching
ques_pos_list = random.sample(range(0, actNoPara, 5),n)
print(ques_pos_list)
for counter,ranNo in enumerate(ques_pos_list):
    #ranNo is the random number
    # Getting text from input file
# for i in range(100,300,5):
    paraCode= InDoc.paragraphs[ranNo]
    InpQues = paraCode.text
    print(InpQues)
    # copying that text into new file
    OutDoc.add_paragraph(f"{counter+1}) {InpQues}")
    # function to randomize options
    optionRandomizer(ranNo)

OutDoc.save("OutDoc.doc")

